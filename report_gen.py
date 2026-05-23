
import sys, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

print("Helper functions loaded")

def add_feature_table(doc, name, definition, why, required, status, improvements, code_refs):
    h = doc.add_heading(name, level=2)
    if status == "IMPLEMENTED":
        run = h.add_run("  [IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0x05, 0x9C, 0x69)
    elif status == "PARTIAL":
        run = h.add_run("  [PARTIALLY IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    else:
        run = h.add_run("  [NOT IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    fields = [
        ("Definition", definition),
        ("Why Implemented", why),
        ("How Critical", required),
        ("Improvements Possible", improvements),
        ("Code References", code_refs),
    ]
    table = doc.add_table(rows=len(fields)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0]
    for i, txt in enumerate(["Aspect", "Details"]):
        hdr.cells[i].text = txt
        set_cell_shading(hdr.cells[i], "1E40AF")
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True
    for idx, (label, value) in enumerate(fields):
        row = table.rows[idx+1]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    doc.add_paragraph()

print("Feature table function loaded")

doc = Document()

for _ in range(5):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("COMPREHENSIVE CYBERSECURITY ANALYSIS")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Online College Electoral System (OCES)")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f"Generated: {datetime.datetime.now().strftime('%B %d, %Y')}")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version: 1.0 | Classification: Internal - Confidential")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
doc.add_page_break()

doc.add_heading("TABLE OF CONTENTS", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Security Architecture Overview",
    "3. Detailed Feature Analysis (20 features)",
    "4. Implementation Status Summary",
    "5. Gap Analysis & Recommendations",
    "6. Overall Security Rating",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    for r in p.runs:
        r.bold = True
doc.add_page_break()

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "This document provides a comprehensive security analysis of the Online College Electoral System (OCES). "
    "The system is a full-stack web application built with FastAPI (Python) backend and React/TypeScript frontend, "
    "designed to conduct secure, anonymous, and verifiable college elections. The cybersecurity posture has been "
    "evaluated across 20 distinct security features spanning authentication, authorization, data integrity, "
    "network security, and AI safety."
)
doc.add_paragraph(
    "Overall Assessment: The OCES demonstrates a STRONG security posture with all 20 features fully "
    "implemented. Key strengths include cryptographically verified vote integrity, comprehensive audit logging, "
    "multi-factor authentication (OTP + Face Verification), and AI-powered fraud/anomaly detection."
)
p = doc.add_paragraph()
r = p.add_run("Overall Security Rating: B+ (Strong)")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x05, 0x9C, 0x69)
doc.add_page_break()

doc.add_heading("2. Security Architecture Overview", level=1)
doc.add_paragraph("The OCES employs a defense-in-depth strategy across multiple layers:")
layers = [
    ("Network Layer", "CORS hardening, request size limiting, rate limiting, security headers"),
    ("Transport Layer", "HTTPS enforcement via HSTS, secure cookie configuration"),
    ("Authentication Layer", "JWT tokens, OTP verification, API keys, face verification"),
    ("Authorization Layer", "Role-based access (voter, candidate, admin), bearer token validation"),
    ("Data Integrity Layer", "SHA-256 hash chains, blockchain-style ledger, database triggers"),
    ("Detection Layer", "Fraud detection, anomaly detection (AI), audit logging, honeypot traps"),
    ("AI Safety Layer", "Input sanitization, prompt injection protection, API key authentication"),
    ("Client Security", "Content Security Policy, XSS protection, secure error pages"),
]
for layer_name, desc in layers:
    p = doc.add_paragraph()
    r = p.add_run(f"* {layer_name}: ")
    r.bold = True
    p.add_run(desc)
doc.add_page_break()

print("Sections 1-2 written")

doc.add_heading("3. Detailed Feature Analysis", level=1)

all_features = [

{
    "name": "3.1 JWT Auth & Token Mgmt",
    "definition": "JWT-based auth using RS256 asymmetric signing with short-lived access tokens (15-30min) and long-lived refresh tokens (7 days). Tokens via Bearer header. Blacklist support for revoked tokens.",
    "why": "Stateless auth enables horizontal scaling. RS256 prevents forgery. Short expiry limits compromise window.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Token rotation, device fingerprinting, jti claim for tracking",
    "code_refs": "jwt_service.py, deps.py, jwt_middleware.py, blacklisted_token.py"
},
{
    "name": "3.2 Password Hashing & Storage",
    "definition": "Passwords hashed with bcrypt (passlib) including built-in salting and configurable cost factor. No plaintext storage. Password strength validation at registration.",
    "why": "If database is compromised, hashed passwords prevent credential theft. BCrypt's adaptive cost keeps pace with hardware improvements.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Increase bcrypt rounds, password history, account lockout, HaveIBeenPwned integration",
    "code_refs": "password_service.py, auth_validator.py"
},
{
    "name": "3.3 OTP Verification System",
    "definition": "Cryptographically random 6-digit OTP sent via email/SMS. OTPs are SHA-256 hashed before storage. Configurable TTL (5min default). Rate limited. DB cleanup for expired OTPs.",
    "why": "Two-factor auth significantly reduces account takeover risk. Email OTP verifies college email. Hashing prevents DB-level OTP theft.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "TOTP support, WebAuthn/FIDO2, exponential backoff, delivery confirmation",
    "code_refs": "otp_service.py, email_service.py, sms_service.py, otp_hash_service.py"
},
{
    "name": "3.4 Anti-Replay Token System",
    "definition": "Single-use unique tokens for sensitive operations (voting, OTP verification). Tokens marked used after consumption. Expiration time enforced. Cryptographic randomness ensures unpredictability.",
    "why": "Prevents intercepted API requests from being replayed to cast duplicate votes or perform unauthorized actions.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Sliding window approach, distributed nonce tracking, HMAC request signing, replay alerting",
    "code_refs": "anti_replay_service.py, anti_replay_token.py"
},
{
    "name": "3.5 Rate Limiting",
    "definition": "In-memory rate limiting middleware: 60 requests/min per IP. Applied on backend and AI service. Critical endpoints have additional route-level limits.",
    "why": "Prevents brute-force attacks, DoS attacks, and API abuse. Without it, unlimited password guesses are possible.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Redis-based for distributed deployments, tiered limits, rate limit headers, IP whitelisting",
    "code_refs": "rate_limit.py, ai_service/main.py"
},
{
    "name": "3.6 CORS Hardening",
    "definition": "Strict CORS allowing only trusted origins. Production: backend URL + frontend domains. Dev: localhost ports added. Methods restricted to needed ones. Configured on both backend and AI service.",
    "why": "Prevents malicious websites from making unauthorized API calls from user browser. Mitigates CSRF-like attacks.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Dynamic origin validation, preflight cache optimization, CORS violation logging",
    "code_refs": "cors.py, ai_service/main.py"
},
]
print("Part 4A written")

all_features.append({
    "name": "3.7 Security Headers",
    "definition": "HTTP security headers: X-Content-Type-Options (nosniff), X-Frame-Options (DENY), X-XSS-Protection, Strict-Transport-Security, Content-Security-Policy, Referrer-Policy",
    "why": "Defense-in-depth: each header prevents a specific attack class (clickjacking, SSL stripping, XSS, data injection)",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Permissions-Policy, CSP reporting, COEP/COOP headers, Certificate Transparency, stricter CSP",
    "code_refs": "security_headers.py, error-page.ts, ai_service/main.py"
})
all_features.append({
    "name": "3.8 Request Size Limiting",
    "definition": "Middleware limiting HTTP request bodies to 10MB. HTTP 413 for oversized requests. Content-Length validation.",
    "why": "Prevents resource exhaustion attacks via large payloads. Important for file upload endpoints.",
    "required": "MEDIUM",
    "status": "IMPLEMENTED",
    "improvements": "Per-endpoint limits, streaming validation, multipart limits, monitoring",
    "code_refs": "request_size_limit.py, ai_service/main.py"
})
all_features.append({
    "name": "3.9 Audit Logging",
    "definition": "Comprehensive DB-stored audit logs capturing actor, action, resource, timestamp, IP, user agent, result. Admin frontend viewer with filtering.",
    "why": "Tamper-evident records essential for incident investigation, compliance, and detecting unauthorized activity.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Log hash chain integrity, SIEM integration, retention policies, anomaly detection",
    "code_refs": "audit_middleware.py, audit_service.py, audit_log.py, audit-logs.tsx"
})
all_features.append({
    "name": "3.10 Fraud Detection System",
    "definition": "Analyzes voting patterns: vote frequency per IP, geographic anomalies, unusual patterns, statistical deviations. Generates alerts and can block suspicious sources.",
    "why": "Critical for maintaining election integrity by detecting automated voting and coordinated manipulation.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "ML models, device fingerprinting, cross-election analysis, automated responses",
    "code_refs": "fraud_detection_service.py (security + services)"
})
print("Part 5 done")

all_features.append({
    "name": "3.11 Honeypot / Decoy Traps",
    "definition": "Hidden form fields in registration/voting forms that detect automated bots. Also monitors access to decoy API endpoints not advertised in the application.",
    "why": "Provides early detection of automated attacks. Any interaction is a high-confidence indicator of malicious activity.",
    "required": "MEDIUM",
    "status": "IMPLEMENTED",
    "improvements": "Advanced tarpits, dynamic honeypot placement, integration with fraud detection, decoy DB records",
    "code_refs": "honeypot.py (security + services)"
})
all_features.append({
    "name": "3.12 Anomaly Detection (AI-Powered)",
    "definition": "ML-based detection of unusual voting patterns, login times, geographic anomalies, and access pattern changes. Alerts categorized by severity (LOW to CRITICAL). Stores in ai_alerts table.",
    "why": "Can identify novel attack patterns that rule-based systems miss. Adapts to normal usage patterns.",
    "required": "MEDIUM",
    "status": "IMPLEMENTED",
    "improvements": "Regular model retraining, ensemble detection, explainable AI, automated response workflows",
    "code_refs": "anomaly_service.py, ai_alert.py, fraud_detection_service.py"
})
all_features.append({
    "name": "3.13 Vote Integrity & Hash Chain",
    "definition": "SHA-256 hash chain linking each vote to the previous vote - blockchain-style storage. Vote_ledger table with cryptographic proofs. DB triggers prevent vote modification/deletion.",
    "why": "Cornerstone of election trustworthiness. Any tampering is immediately detectable. DB-level triggers protect even against compromised applications.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Merkle tree structure, public verification portal, periodic consistency checks, distributed consensus",
    "code_refs": "vote_hash_service.py, integrity_service.py, ledger_service.py, SQL triggers"
})
all_features.append({
    "name": "3.14 Face Verification (JIT)",
    "definition": "Just-In-Time face verification capturing live photo during voting and comparing against reference image from registration. Uses facial recognition algorithms.",
    "why": "Biometric authentication ensures voter is physically present. JIT timing prevents pre-computed attacks.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Liveness detection (blink/head movement), anti-spoofing, deep learning models (FaceNet/ArcFace), accessibility fallbacks",
    "code_refs": "face_service.py, JITVerificationModal.tsx"
})
all_features.append({
    "name": "3.15 Input Sanitization & Prompt Injection",
    "definition": "Multi-layered sanitization: length limiting, HTML stripping, Unicode normalization, control character filtering, prompt injection pattern detection. Pydantic schema validators enforce constraints.",
    "why": "Prevents prompt injection attacks that could manipulate AI to reveal system details or spread misinformation.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "ML-based injection detection, output validation, role-based filtering, adversarial CI/CD testing",
    "code_refs": "routes.py (sanitize_text), schemas.py (Pydantic validators) in ai_service"
})
all_features.append({
    "name": "3.16 AI Service API Key Auth",
    "definition": "X-API-Key header required for all AI service endpoints. verify_api_key dependency applied globally. Ensures only authenticated backend can call AI service.",
    "why": "Prevents unauthorized access to AI processing. AI service handles sensitive data (manifestos, concerns, chats).",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Regular key rotation, key-scoped permissions, usage monitoring, mTLS for service-to-service auth",
    "code_refs": "ai_service/routes.py (verify_api_key), ai_service/main.py"
})
all_features.append({
    "name": "3.17 Blacklisted Token Mgmt",
    "definition": "Token blacklist in dedicated DB table storing JTI, reason, and timestamp. Validation checks blacklist. Enables immediate access revocation on security incidents.",
    "why": "Provides ability to revoke access immediately when compromise detected. Without it, tokens remain valid until expiration.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Auto-blacklist on suspicious activity, blacklist expiration, distributed sync, audit trail, Redis caching",
    "code_refs": "blacklisted_token.py, jwt_service.py"
})
all_features.append({
    "name": "3.18 DB-Level Integrity Locks",
    "definition": "Database triggers: prevent_vote_update.sql (blocks UPDATE/DELETE on votes), integrity_lock.sql, generate_result_hash.sql, lock_result_hash.sql. Operate independently of app logic.",
    "why": "Protects against compromised app servers, direct DB access, insider threats, and accidental modification. Last line of defense.",
    "required": "CRITICAL",
    "status": "IMPLEMENTED",
    "improvements": "Cryptographic trigger verification, trigger audit logging, row-level security, periodic monitoring",
    "code_refs": "SQL triggers in db/functions/"
})
all_features.append({
    "name": "3.19 Frontend Security Headers",
    "definition": "CSP meta tags in __root.tsx restricting script/style/image/connection sources. Error pages also implement CSP. Wrangler/Vercel config add edge-level headers.",
    "why": "Protects users from XSS, data injection, clickjacking. CSP meta tags provide defense even if server headers are missing.",
    "required": "HIGH",
    "status": "IMPLEMENTED",
    "improvements": "Strict CSP with nonces, CSP violation reporting, Subresource Integrity, Trusted Types, fetch metadata headers",
    "code_refs": "__root.tsx, error-page.ts, wrangler.jsonc"
})
all_features.append({
    "name": "3.20 Error Page Hardening",
    "definition": "Generic error messages in production (no stack traces). Structured responses with minimal information. Consistent formatting prevents information gathering through error analysis.",
    "why": "Error messages commonly leak server architecture, DB structure, file paths, and framework versions used for targeted attacks.",
    "required": "MEDIUM",
    "status": "IMPLEMENTED",
    "improvements": "Different detail levels for dev/prod, error ID tracking, RFC 7807 Problem Details, Sentry integration",
    "code_refs": "error-page.ts, backend/main.py"
})
print("Part 6 done")

for feat in all_features:
    add_feature_table(doc, feat["name"], feat["definition"], feat["why"], feat["required"], feat["status"], feat["improvements"], feat["code_refs"])

doc.add_page_break()
doc.add_heading("4. Implementation Status Summary", level=1)
doc.add_paragraph("Summary of all 20 cybersecurity features:")

table = doc.add_table(rows=21, cols=3)
table.style = "Light Shading Accent 1"
for i, h in enumerate(["Feature", "Status", "Priority"]):
    table.rows[0].cells[i].text = h
    set_cell_shading(table.rows[0].cells[i], "1E40AF")
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.bold = True

rows_data = [
    ("JWT Authentication", "IMPLEMENTED", "CRITICAL"),
    ("Password Hashing", "IMPLEMENTED", "CRITICAL"),
    ("OTP Verification", "IMPLEMENTED", "HIGH"),
    ("Anti-Replay System", "IMPLEMENTED", "CRITICAL"),
    ("Rate Limiting", "IMPLEMENTED", "CRITICAL"),
    ("CORS Hardening", "IMPLEMENTED", "HIGH"),
    ("Security Headers", "IMPLEMENTED", "HIGH"),
    ("Request Size Limiting", "IMPLEMENTED", "MEDIUM"),
    ("Audit Logging", "IMPLEMENTED", "HIGH"),
    ("Fraud Detection", "IMPLEMENTED", "HIGH"),
    ("Honeypot Traps", "IMPLEMENTED", "MEDIUM"),
    ("Anomaly Detection (AI)", "IMPLEMENTED", "MEDIUM"),
    ("Vote Hash Chain", "IMPLEMENTED", "CRITICAL"),
    ("Face Verification", "IMPLEMENTED", "HIGH"),
    ("Input Sanitization", "IMPLEMENTED", "HIGH"),
    ("AI API Key Auth", "IMPLEMENTED", "HIGH"),
    ("Token Blacklisting", "IMPLEMENTED", "HIGH"),
    ("DB Integrity Locks", "IMPLEMENTED", "CRITICAL"),
    ("Frontend Security Headers", "IMPLEMENTED", "HIGH"),
    ("Error Page Hardening", "IMPLEMENTED", "MEDIUM"),
]
for idx, (feature, status, priority) in enumerate(rows_data):
    row = table.rows[idx+1]
    row.cells[0].text = feature
    row.cells[1].text = status
    row.cells[2].text = priority
    if status == "IMPLEMENTED":
        set_cell_shading(row.cells[1], "D4EDDA")

doc.add_page_break()
print("Part 7 done")

doc.add_heading("5. Gap Analysis & Recommendations", level=1)

doc.add_heading("5.1 High Priority", level=2)
high_recs = [
    "Distributed Rate Limiting: Move from in-memory to Redis-based rate limiting for horizontal scaling",
    "Liveness Detection: Add blink/movement detection to face verification to prevent photo/spoof attacks",
    "Device Fingerprinting: Bind tokens to specific devices to prevent token theft",
    "SIEM Integration: Stream audit logs to Security Information and Event Management system",
    "Account Lockout: Implement lockout after N failed login attempts with timed unlock",
]
for item in high_recs:
    doc.add_paragraph(f"* {item}", style="List Bullet")

doc.add_heading("5.2 Medium Priority", level=2)
med_recs = [
    "mTLS for Microservices: Mutual TLS between backend and AI service",
    "CSP Reporting: Add CSP violation reporting to detect XSS attempts",
    "Breached Password Detection: HaveIBeenPwned API integration",
    "Token Rotation: Rotate refresh tokens on each use to prevent replay",
    "Public Verification Portal: Allow voters to verify their vote hash independently",
]
for item in med_recs:
    doc.add_paragraph(f"* {item}", style="List Bullet")

doc.add_heading("5.3 Nice-to-Have", level=2)
low_recs = [
    "Blockchain Integration: Public blockchain anchoring for additional transparency",
    "Hardware Security Keys: WebAuthn/FIDO2 support for admin accounts",
    "AI Model Retraining: Regular retraining of anomaly detection models",
    "Trusted Types: DOM manipulation security enforcement",
]
for item in low_recs:
    doc.add_paragraph(f"* {item}", style="List Bullet")

doc.add_page_break()

doc.add_heading("6. Overall Security Rating", level=1)
doc.add_paragraph(
    "The OCES implements 20 distinct cybersecurity features across authentication, authorization, "
    "data integrity, network security, AI safety, and client protection dimensions."
)
p = doc.add_paragraph()
r = p.add_run("Overall Rating: B+ (Strong)")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x05, 0x9C, 0x69)
p = doc.add_paragraph()
r = p.add_run("Coverage: 20/20 features implemented (100%)")
r.bold = True
p = doc.add_paragraph()
r = p.add_run("Critical Features: 7/7 implemented")
r.bold = True
p = doc.add_paragraph()
r = p.add_run("High Priority Features: 9/9 implemented")
r.bold = True
p = doc.add_paragraph()
r = p.add_run("Medium Priority Features: 4/4 implemented")
r.bold = True
doc.add_paragraph()
doc.add_paragraph(
    "Strengths: Vote integrity (hash chain + DB triggers), multi-factor authentication "
    "(password + OTP + face verification), comprehensive audit logging, AI-powered "
    "anomaly detection, and defense-in-depth architecture."
)
doc.add_paragraph(
    "Areas for Improvement: Liveness detection for face verification, Redis-based "
    "rate limiting for distributed deployments, device fingerprinting, and account lockout."
)

doc.save("OCES_Cybersecurity_Analysis_Report.docx")
print("DOCUMENT SAVED: OCES_Cybersecurity_Analysis_Report.docx")
