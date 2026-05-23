
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_feature_table(doc, name, definition, why, required, status, improvements, code_refs):
    h = doc.add_heading(name, level=2)
    if status == "IMPLEMENTED":
        run = h.add_run("  [IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0x05, 0x9C, 0x69)
    elif status == "PARTIAL":
        run = h.add_run("  [PARTIALLY IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    else:
        run = h.add_run("  [NOT IMPLEMENTED]")
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    
    fields = [
        ("Definition", definition),
        ("Why Implemented", why),
        ("How Critical", required),
        ("Improvements Possible", improvements),
        ("Code References", code_refs),
    ]
    table = doc.add_table(rows=len(fields)+1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0]
    for i, txt in enumerate(["Aspect", "Details"]):
        hdr.cells[i].text = txt
        set_cell_shading(hdr.cells[i], "1E40AF")
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True
    for idx, (label, value) in enumerate(fields):
        row = table.rows[idx+1]
        row.cells[0].text = label
        row.cells[1].text = str(value)
    doc.add_paragraph()

print("Functions defined OK")
