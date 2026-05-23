
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# Title page
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('COMPREHENSIVE CYBERSECURITY ANALYSIS')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Online College Electoral System (OCES)')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

doc.add_paragraph()
doc.add_paragraph()

doc.add_heading('TABLE OF CONTENTS', level=1)
toc = ['1. Executive Summary','2. System Architecture Overview','3. Implemented Features (Detailed Analysis)','4. Features Yet to Implement (Gap Analysis)','5. Overall Assessment Score','6. Prioritized Remediation Roadmap','7. Final Recommendations','8. Appendix']
for t in toc:
    doc.add_paragraph(t)

doc.add_page_break()

# Executive Summary
doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
doc.add_paragraph(
    'The Online College Electoral System (OCES) has been designed with a defense-in-depth security architecture. '
    'This report provides an exhaustive analysis of all cybersecurity features currently implemented, their definitions, '
    'rationales, criticality ratings, and actionable improvement recommendations.')
doc.add_paragraph(
    'Overall, OCES demonstrates a robust security posture with 17 major cybersecurity features '
    'implemented across authentication, authorization, network security, cryptography, monitoring, and '
    'physical security domains. The system implements bcrypt password hashing, JWT-based session management, '
    'multi-factor authentication via OTP, rate limiting, audit logging, and cryptographic hash chain for vote integrity.')

doc.add_paragraph()
r = doc.add_paragraph().add_run('Overall Security Rating: B+ (Strong) | Score: 83/100')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x05, 0x9C, 0x69)

doc.add_paragraph(
    'The system has a well-architected security foundation, with notable strengths in cryptographic '
    'vote integrity, anti-fraud measures, and comprehensive audit logging. Key gaps include the absence '
    'of a Web Application Firewall (WAF), database encryption at rest, and formal penetration testing.')

doc.save('OCES_Cybersecurity_Analysis_Report.docx')
print('BASE_DONE')
